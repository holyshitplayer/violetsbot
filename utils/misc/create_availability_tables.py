import docx
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm

from utils.db_api import Database


def availability_table_content(db: Database, mode: str, list_title: str):
    if mode == "orders_in_boxes":
        list_content = db.list_content_kids(list_title)
    else:
        list_content = db.list_content_orders(list_title)

    table_content = {"title": list_title, "content": []}

    for sort in list_content:
        sort_id = sort[0]
        sort_count = sort[1]

        if mode == "orders_in_boxes":
            sort_count_in_other_table = db.sort_count_orders(sort_id)
        else:
            sort_count_in_other_table = db.sort_count_kids(sort_id)

        if sort_count_in_other_table > 0:
            sort_name = db.select_sort(id=sort_id)[1]
            if mode == "orders_in_boxes":
                sorts_in_other_table = db.select_sort_in_orders(sort_id)
            else:
                sorts_in_other_table = db.select_sort_in_kids(sort_id)
            table_content["content"].append({"sort_name": sort_name, "sort_count": sort_count, "sort_in_other_table": sorts_in_other_table})
        else:
            continue

    return table_content


def create_availability_table(doc: docx.Document, mode: str, list_title: str, table_content):
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(8)
    title_content = title.add_run(list_title)
    title_content.font.size = Pt(14)
    title_content.font.bold = True

    table = doc.add_table(rows=1, cols=3, style="Table Grid")
    table.cell(0, 0).width = Cm(4)
    table.cell(0, 1).width = Cm(2)
    table.cell(0, 2).width = Cm(12.5)
    table_title_cells = table.rows[0].cells
    table_title_cells[0].text = "Сорт"
    table_title_cells[1].text = "Кол-во"
    table_title_cells[2].text = "Заказ № (кол-во)" if mode == "orders_in_boxes" else "Ящик № (кол-во)"

    for sort in table_content["content"]:
        sort_list = ""
        sort_list += "\n".join([
            f"{list_name} ({list_count} шт.)" for list_name, list_count in sort["sort_in_other_table"]
        ])

        row_cells = table.add_row().cells
        row_cells[0].width = Cm(10)
        row_cells[1].width = Cm(2)
        row_cells[2].width = Cm(4)
        row_cells[0].text = sort["sort_name"]
        row_cells[1].text = str(sort["sort_count"])
        row_cells[2].text = sort_list

    doc.add_page_break()
    return doc


def availability_table_document(db: Database, mode: str, list_title: str):
    doc = Document("utils/misc/template.docx")
    style = doc.styles["Normal"]
    style.paragraph_format.space_after = Pt(0)
    style_font = style.font
    style_font.name = "Proxima Nova Rg"
    style_font.size = Pt(12)

    table_content = availability_table_content(db, mode, list_title)

    doc = create_availability_table(doc, mode, list_title, table_content)

    doc_path = f"data/files/{list_title} - таблица наличия.docx"
    doc.save(doc_path)

    return doc_path


def availability_tables_document(db: Database, mode: str):
    doc = Document("utils/misc/template.docx")
    style = doc.styles["Normal"]
    style.paragraph_format.space_after = Pt(0)
    style_font = style.font
    style_font.name = "Proxima Nova Rg"
    style_font.size = Pt(12)

    if mode == "orders_in_boxes":
        list_titles = db.list_titles_kids()
    else:
        list_titles = db.list_titles_orders()

    for list_title_item in list_titles:
        list_title = list_title_item[0]
        table_content = availability_table_content(db, mode, list_title)
        if len(table_content["content"]) == 0:
            continue
        doc = create_availability_table(doc, mode, list_title, table_content)

    doc_path = f"data/files/Таблицы наличия.docx"
    doc.save(doc_path)

    return doc_path
