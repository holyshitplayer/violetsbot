from docx import Document
from docx.shared import Pt

from utils.db_api import Database


def export_list_to_docx(db: Database, table_name: str, list_title: str):
    doc = Document("utils/misc/template.docx")
    style = doc.styles["Normal"]
    style.paragraph_format.space_after = Pt(3)
    style.font.name = "Proxima Nova Rg"
    style.font.size = Pt(12)
    style.font.bold = True
    title = doc.add_paragraph().add_run(list_title.upper())
    title.font.size = Pt(20)

    list_content = None
    if table_name == "kids":
        list_content = db.list_content_kids(list_title)
    elif table_name == "starters":
        list_content = db.list_content_starters(list_title)
    elif table_name == "orders":
        list_content = db.list_content_orders(list_title)

    for i, sort in enumerate(list_content):
        sort_name = db.select_sort(id=sort[0])[1]
        if sort[1] == 1:
            doc.add_paragraph(f"{i + 1}) {sort_name}")
        else:
            doc.add_paragraph(f"{i + 1}) {sort_name} ({sort[1]})")

    doc_path = f"data/files/{list_title}.docx"
    doc.save(doc_path)
    return doc_path
